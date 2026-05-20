import os
import shutil
import PyPDF2
import docx
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from openai import OpenAI

try:
    from fpdf import FPDF
    PDF_GENERATION_AVAILABLE = True
except ImportError:
    FPDF = None
    PDF_GENERATION_AVAILABLE = False

from dotenv import load_dotenv, find_dotenv
import json
import traceback
from utils.api_config import get_llm_provider
from utils.openai_patch import patch_openai

# Apply the OpenAI patch
patch_openai()

# Load environment variables from project root or fallback to static/.env
env_loaded = load_dotenv(find_dotenv())
if not env_loaded:
    fallback_env = os.path.join(os.path.dirname(__file__), '..', 'static', '.env')
    fallback_env = os.path.normpath(fallback_env)
    if os.path.exists(fallback_env):
        load_dotenv(fallback_env)

# Force reload .env
from dotenv import load_dotenv
load_dotenv(override=True)  # Force override

# Get the LLM provider
provider = get_llm_provider()
if provider:
    provider_name = os.getenv("LLM_PROVIDER", "OpenAI").capitalize()
    print(f"✅ Resume analyzer: {provider_name} provider initialized successfully")
    print(f"   Provider object: {provider}")
else:
    print("❌ Resume analyzer: Using basic evaluation (LLM provider not available)")   
# Download NLTK data if not present
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    try:
        nltk.download('punkt')
    except:
        print("Warning: Could not download NLTK punkt tokenizer")

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    try:
        nltk.download('stopwords')
    except:
        print("Warning: Could not download NLTK stopwords")


# Common ATS keywords by job role
JOB_KEYWORDS = {
    "Software Developer": [
        "python", "javascript", "java", "c++", "react", "node.js", "database", 
        "api", "cloud", "aws", "azure", "agile", "scrum", "git", "algorithms",
        "data structures", "backend", "frontend", "full-stack", "devops"
    ],
    "Data Scientist": [
        "python", "r", "machine learning", "sql", "statistics", "data analysis",
        "tensorflow", "pytorch", "pandas", "numpy", "visualization", "big data",
        "hadoop", "spark", "ai", "algorithms", "modeling", "nlp", "data mining"
    ],
    "Project Manager": [
        "project management", "agile", "scrum", "kanban", "leadership", "budget",
        "stakeholder", "risk management", "pmp", "planning", "reporting", "kpi",
        "team lead", "coordination", "strategy", "delivery", "timeline", "milestone"
    ],
    "General": [
        "leadership", "communication", "teamwork", "problem-solving", "analytical",
        "detail-oriented", "time management", "project", "innovation", "collaboration",
        "customer service", "organization", "adaptability", "creativity", "initiative"
    ]
}

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text()
        if not text.strip():
            print(f"Warning: Extracted empty text from PDF: {file_path}")
            return "Empty PDF content. Please check your file is not corrupted or password protected."
        return text
    except Exception as e:
        print(f"Error reading PDF: {str(e)}")
        traceback.print_exc()
        raise Exception(f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = " ".join([paragraph.text for paragraph in doc.paragraphs])
        if not text.strip():
            print(f"Warning: Extracted empty text from DOCX: {file_path}")
            return "Empty DOCX content. Please check your file."
        return text
    except Exception as e:
        print(f"Error reading DOCX: {str(e)}")
        traceback.print_exc()
        raise Exception(f"Error reading DOCX: {str(e)}")

def extract_text_from_file(file_path):
    """Extract text from supported file formats."""
    _, file_extension = os.path.splitext(file_path)
    
    print(f"Processing file: {file_path} with extension: {file_extension}")
    
    if file_extension.lower() == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension.lower() == '.docx':
        return extract_text_from_docx(file_path)
    else:
        error_msg = f"Unsupported file format: {file_extension}"
        print(error_msg)
        raise ValueError(error_msg)

def analyze_resume(file_path, job_role="General"):
    """
    Analyze a resume and provide feedback.
    
    Args:
        file_path: Path to the resume file
        job_role: Target job role (default: General)
        
    Returns:
        Dictionary with analysis results
    """
    try:
        print(f"Analyzing resume: {file_path} for job role: {job_role}")
        
        # Verify file exists
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            print(error_msg)
            raise FileNotFoundError(error_msg)
            
        # Extract text from resume
        resume_text = extract_text_from_file(file_path)
        
        # Basic analysis
        analysis = {}
        
        # Check length
        try:
            words = word_tokenize(resume_text)
        except Exception as e:
            print(f"NLTK tokenization failed: {e}, falling back to simple split")
            # Fallback if NLTK fails
            words = resume_text.split()
        
        analysis['word_count'] = len(words)
        
        # Remove stopwords for keyword analysis
        try:
            stop_words = set(stopwords.words('english'))
            words_no_stop = [word.lower() for word in words if word.lower() not in stop_words and word.isalpha()]
        except Exception as e:
            print(f"Stopwords removal failed: {e}, falling back to simple filtering")
            # Fallback if NLTK fails
            words_no_stop = [word.lower() for word in words if word.isalpha()]
        
        # Get relevant keywords for job role
        target_keywords = JOB_KEYWORDS.get(job_role, JOB_KEYWORDS["General"])
        
        # Check for keyword matches
        found_keywords = [keyword for keyword in target_keywords if keyword in resume_text.lower()]
        missing_keywords = [keyword for keyword in target_keywords if keyword not in resume_text.lower()]
        
        analysis['keywords_found'] = found_keywords
        analysis['keywords_missing'] = missing_keywords
        analysis['keyword_match_percent'] = round((len(found_keywords) / len(target_keywords)) * 100, 2)
        
        # Use AI to generate detailed feedback if provider is available
        if provider:
            try:
                print("Attempting to get AI feedback for resume")
                ai_feedback = get_ai_resume_feedback(resume_text, job_role, found_keywords, missing_keywords)
                analysis.update(ai_feedback)
                print("Successfully obtained AI feedback")
            except Exception as e:
                print(f"AI feedback failed: {e}")
                traceback.print_exc()
                analysis['ai_error'] = str(e)
                # Provide basic feedback if AI fails
                analysis['feedback'] = get_basic_feedback(job_role, found_keywords, missing_keywords)
                analysis['score'] = 65  # Default score
        else:
            # Provide basic feedback if LLM provider is not available
            print("Using basic feedback (LLM provider not available)")
            analysis['feedback'] = get_basic_feedback(job_role, found_keywords, missing_keywords)
            analysis['score'] = 65  # Default score
        
        return analysis
    except Exception as e:
        print(f"Resume analysis failed with error: {e}")
        traceback.print_exc()
        raise

def get_basic_feedback(job_role, found_keywords, missing_keywords):
    """Provide basic feedback when AI is not available."""
    try:
        total_keywords = len(found_keywords) + len(missing_keywords)
        match_percent = round((len(found_keywords) / total_keywords) * 100, 1) if total_keywords > 0 else 0
        
        # Create bullet point lists
        found_bullets = '\n'.join([f"- {keyword}" for keyword in found_keywords[:5]])
        missing_bullets = '\n'.join([f"- {keyword}" for keyword in missing_keywords[:5]])
        
        return {
            "structure": "Resume structure appears standard. Consider the following guidelines:\n- Ensure you have clear sections for experience, education, and skills\n- Use consistent formatting for dates and headings\n- Include a professional summary at the top\n- Keep the layout clean and easy to scan",
            "content": f"Your resume contains {len(found_keywords)} relevant keywords for {job_role} position. Top keywords found:\n{found_bullets}",
            "improvements": f"Consider adding these missing keywords to improve ATS compatibility:\n{missing_bullets}\n- Add specific achievements with metrics\n- Tailor your experience to the {job_role} position\n- Use action verbs to start bullet points",
            "summary": f"Your resume has a {match_percent}% keyword match for {job_role} positions. Focus on adding missing keywords and quantifying your achievements."
        }
    except Exception as e:
        print(f"Error generating basic feedback: {e}")
        return {
            "structure": "Unable to analyze structure due to an error.",
            "content": "Unable to analyze content due to an error.",
            "improvements": "Unable to provide improvements due to an error.",
            "summary": "An error occurred during resume analysis."
        }

def get_ai_resume_feedback(resume_text, job_role, found_keywords, missing_keywords):
    """
    Use LLM provider to generate detailed resume feedback.
    """
    # Get provider FRESH each time
    from utils.api_config import get_llm_provider
    provider = get_llm_provider()
    
    if not provider:
        raise Exception("LLM provider not available")
    
    # Limit text length to avoid token limit
    truncated_text = resume_text[:3000] if len(resume_text) > 3000 else resume_text
    
    prompt = f"""Analyze this resume for an applicant seeking a {job_role} position.

Resume text:
{truncated_text}

The resume contains these relevant keywords: {', '.join(found_keywords[:10])}
It's missing these potentially important keywords: {', '.join(missing_keywords[:10])}

Provide your response as a valid JSON object with these EXACT keys:
- structure (string): assessment of resume structure and formatting
- content (string): feedback on content quality and relevance
- improvements (string): specific recommendations for improvement
- score (integer): a number from 0-100
- summary (string): brief summary of strengths and weaknesses

Return ONLY the JSON object. No explanations before or after. Use this exact format:
{{"structure": "...", "content": "...", "improvements": "...", "score": 75, "summary": "..."}}"""

    try:
        provider_name = os.getenv("LLM_PROVIDER", "LLM").upper()
        print(f"🤖 Sending request to {provider_name} API...")
        
        feedback_text = provider.chat_completion(
            messages=[
                {"role": "system", "content": "You are an expert ATS resume analyzer and HR professional. Always respond with valid JSON only. Never include markdown, code blocks, or extra text."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        print(f"📥 Received response from {provider_name} API: {len(feedback_text)} chars")
        
        # Clean the response
        feedback_text = feedback_text.strip()
        
        # Remove markdown code blocks
        if feedback_text.startswith("```json"):
            feedback_text = feedback_text[7:]
        if feedback_text.startswith("```"):
            feedback_text = feedback_text[3:]
        if feedback_text.endswith("```"):
            feedback_text = feedback_text[:-3]
        feedback_text = feedback_text.strip()
        
        # Extract JSON object using brace counting
        brace_count = 0
        start_idx = -1
        end_idx = -1
        
        for i, char in enumerate(feedback_text):
            if char == '{':
                if brace_count == 0:
                    start_idx = i
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0 and start_idx != -1:
                    end_idx = i
                    break
        
        if start_idx != -1 and end_idx != -1:
            json_str = feedback_text[start_idx:end_idx+1]
            print(f"📋 Extracted JSON: {json_str[:200]}...")
            feedback_data = json.loads(json_str)
        else:
            # Try regex as fallback
            json_match = re.search(r'\{[\s\S]*\}', feedback_text)
            if json_match:
                json_str = json_match.group()
                feedback_data = json.loads(json_str)
            else:
                raise ValueError("No valid JSON found in response")
        
        # Ensure all required keys exist
        required_keys = ['structure', 'content', 'improvements', 'score', 'summary']
        for key in required_keys:
            if key not in feedback_data:
                if key == 'score':
                    feedback_data[key] = 65
                else:
                    feedback_data[key] = f"Information about {key} not available in AI response."
        
        # Ensure score is a number
        if not isinstance(feedback_data['score'], (int, float)):
            try:
                feedback_data['score'] = int(feedback_data['score'])
            except:
                feedback_data['score'] = 65
        
        # Clamp score to 0-100
        feedback_data['score'] = max(0, min(100, feedback_data['score']))
        
        print(f"✅ Successfully parsed AI feedback. Score: {feedback_data['score']}/100")
        
        return {
            'feedback': feedback_data,
            'score': feedback_data['score']
        }
        
    except json.JSONDecodeError as e:
        print(f"⚠️ JSON parse error: {e}")
        print(f"Raw response: {feedback_text[:500]}")
        # Fallback to basic feedback
        basic = get_basic_feedback(job_role, found_keywords, missing_keywords)
        return {
            'feedback': basic,
            'score': 65
        }
        
    except Exception as e:
        print(f"❌ AI API call failed: {e}")
        traceback.print_exc()
        basic = get_basic_feedback(job_role, found_keywords, missing_keywords)
        return {
            'feedback': basic,
            'score': 65
        }

def generate_resume_update_summary(analysis, job_role):
    """Generate a brief update summary for the resume."""
    feedback = analysis.get('feedback', {}) if isinstance(analysis, dict) else {}
    improvements = feedback.get('improvements', '')
    summary_lines = [
        f"Updated resume for: {job_role}",
        f"Score: {analysis.get('score', 0) if isinstance(analysis, dict) else 0}/100",
        "",
        "AI Improvement Suggestions:",
    ]

    if improvements:
        summary_lines.extend(improvements.split('\n'))
    else:
        summary_lines.append('- Add stronger section headings and quantify accomplishments.')
        summary_lines.append('- Include more job-specific keywords throughout the resume.')
        summary_lines.append('- Use concise bullet points and action verbs in experience descriptions.')

    found = analysis.get('keywords_found', []) if isinstance(analysis, dict) else []
    missing = analysis.get('keywords_missing', []) if isinstance(analysis, dict) else []

    if found:
        summary_lines.extend(['', 'Keywords found:'] + [f'- {keyword}' for keyword in found[:8]])

    if missing:
        summary_lines.extend(['', 'Keywords missing:'] + [f'- {keyword}' for keyword in missing[:8]])

    return '\n'.join(summary_lines)


def create_updated_resume(file_path, job_role, analysis, original_filename):
    """Create an updated resume file based on analysis feedback."""
    output_dir = os.path.dirname(file_path)
    base, extension = os.path.splitext(original_filename)
    new_filename = f"{base}_new{extension}"
    new_path = os.path.join(output_dir, new_filename)

    if extension.lower() == '.docx':
        return create_updated_docx(file_path, job_role, analysis, new_path, new_filename)
    elif extension.lower() == '.pdf':
        return create_updated_pdf(file_path, job_role, analysis, new_path, new_filename)
    else:
        raise ValueError('Unsupported file format for resume update generation.')


def create_updated_docx(file_path, job_role, analysis, new_path, new_filename):
    """Generate a new DOCX resume with an appended improvement summary."""
    try:
        document = docx.Document(file_path)
    except Exception as e:
        print(f"Could not open DOCX file for updating: {e}")
        document = docx.Document()
        document.add_paragraph('Original resume content could not be preserved. This document contains updated resume notes.')

    document.add_page_break()
    document.add_heading('AI Resume Improvement Summary', level=1)
    summary_text = generate_resume_update_summary(analysis, job_role)

    for line in summary_text.split('\n'):
        if not line.strip():
            continue
        if line.strip().startswith('- '):
            document.add_paragraph(line.strip()[2:], style='List Bullet')
        else:
            document.add_paragraph(line)

    document.save(new_path)
    return new_path, new_filename


def create_updated_pdf(file_path, job_role, analysis, new_path, new_filename):
    """Generate a simple updated PDF if possible, otherwise copy the original PDF."""
    if PDF_GENERATION_AVAILABLE and FPDF is not None:
        try:
            resume_text = extract_text_from_file(file_path)
            summary_text = generate_resume_update_summary(analysis, job_role)

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font('Arial', size=11)
            pdf.multi_cell(0, 7, f"Updated Resume for {job_role}\n\n{summary_text}\n\nOriginal Resume Text:\n{resume_text[:14000]}")
            pdf.output(new_path)
            return new_path, new_filename
        except Exception as e:
            print(f"Failed to generate updated PDF: {e}")

    try:
        shutil.copy2(file_path, new_path)
        return new_path, new_filename
    except Exception as e:
        print(f"Failed to copy updated PDF fallback file: {e}")
        raise

def optimize_resume_with_ai(resume_text, job_role):
    """
    Use AI to rewrite and optimize the resume.
    """
    from utils.api_config import get_llm_provider
    provider = get_llm_provider()
    
    if not provider:
        return None, "AI provider not available"
    
    truncated_text = resume_text[:4000] if len(resume_text) > 4000 else resume_text
    
    prompt = f"""Rewrite and optimize this resume for a {job_role} position.

Rules:
1. Replace weak verbs (helped→spearheaded, did→executed, worked→developed)
2. Add [XX%] placeholders for missing metrics
3. Use STAR format (Situation, Task, Action, Result)
4. Keep ALL factual information (dates, companies, titles)
5. Return ONLY the rewritten resume text, no explanations

Resume:
{truncated_text}

Rewritten Resume:"""

    try:
        print("🤖 Generating optimized resume...")
        
        optimized_text = provider.chat_completion(
            messages=[
                {"role": "system", "content": "You are an expert resume writer. Rewrite resumes to be more professional, impactful, and ATS-friendly."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        print(f"✅ Resume optimization complete. Length: {len(optimized_text)} chars")
        return optimized_text, None
        
    except Exception as e:
        print(f"❌ Optimization failed: {e}")
        return None, str(e)